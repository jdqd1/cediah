from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.shared import RGBColor

root = Path.cwd()
source = root / "CEDIAH Web Status.docx"
output = root / "tmp" / "codex-content-status" / "CEDIAH Web Status.updated.docx"
doc = Document(source)

old_status = "[ ] Todavía faltan contenido académico autorizado, roles/matrículas, progreso persistente, recursos, evaluaciones, certificados, flujo editorial y pruebas negativas de acceso."
new_status = "[~] El estudio de contenido dinámico, los roles editoriales, el catálogo y las pruebas negativas están implementados localmente; la migración remota, el QA autenticado con Storage real, las matrículas, los certificados y el despliegue de esta entrega siguen pendientes."
matched = [paragraph for paragraph in doc.paragraphs if paragraph.text == old_status]
if len(matched) != 1:
    raise RuntimeError(f"Expected one technical status paragraph, found {len(matched)}")
matched[0].text = new_status
matched[0].style = "Checklist"

if len(doc.tables) < 54:
    raise RuntimeError(f"Unexpected table count: {len(doc.tables)}")
template = doc.tables[48]
checkpoint = next(
    (table for table in doc.tables if table.cell(0, 0).text.startswith("Siguiente punto de control")),
    None,
)
if checkpoint is None:
    raise RuntimeError("Checkpoint table not found")

new_tbl_xml = deepcopy(template._tbl)
checkpoint._tbl.addprevious(new_tbl_xml)
change_table = Table(new_tbl_xml, doc._body)
heading = doc.add_paragraph("Ficha de cambio - CHG-006", style="Heading 2")
new_tbl_xml.addprevious(heading._p)
rows = [
    ("Campo", "Contenido"),
    ("ID / fecha / autor", "CHG-006 / 10/08/2026 / Codex con José"),
    ("Fase y objetivo", "Fases 2 y 3 - sustituir las superficies estáticas de contenido y habilitar un flujo editorial RBAC funcional y comprobable para administradores y miembros autorizados."),
    ("Qué cambió", "Dashboard, guías, biblioteca y páginas de detalle consumen contenido publicado. El panel protegido permite crear y editar videos, guías, cuestionarios, flashcards y temas; adjuntar MP4/WebM/MOV o PDF; revisar, solicitar cambios, aprobar, publicar y archivar."),
    ("Cómo se implementó", "Contratos Zod compartidos; API Fastify como límite de autorización; BFF de Next para conservar el token en servidor; contenido estructurado por tipo; estados draft -> in_review -> changes_requested/approved -> published -> archived; catálogos y actividades interactivas renderizados desde la API."),
    ("Archivos / migraciones / servicios", "packages/contracts/src/index.ts; apps/api/src/content-authorization.ts, providers/supabase-content.ts y rutas /v1/content + /v1/editor; apps/web/src/app/biblioteca, guias, dashboard, panel/contenido y API BFF; componentes content-library, content-detail y content-studio; migración supabase/migrations/20260810211907_add_dynamic_content_studio.sql; bucket privado previsto content-assets."),
    ("Seguridad y privacidad", "La service key permanece solo en Fastify. Cada mutación valida sesión, roles y propiedad; colaboradores solo gestionan sus borradores y coordinación/administración publica. RLS y grants niegan acceso directo del navegador; bucket privado, carga firmada, ruta UUID, allowlist MIME y límite de 500 MB. Finalización verifica tamaño/MIME real en Storage; conflictos usan updatedAt y las acciones sensibles se auditan."),
    ("Validación automática", "PASS final: pnpm lint; pnpm typecheck; pnpm test (4 archivos, 20/20); build de contratos y API; build web Next 16.2.12 con 25 páginas/rutas; git diff --check. Supabase CLI no pudo ejecutar db lint porque no existe Postgres local en 127.0.0.1:54322; la migración no se aplicó a una base remota."),
    ("Validación en navegador", "VER-018 PASS escritorio: dashboard, guías y biblioteca dinámicos, estados indisponible/vacío y navegación sin overlay ni overflow. VER-019 PASS 390 x 844: filtros de biblioteca en cuadrícula, selector fluido y cero overflow. VER-020 PASS en fixture local retirado: estudio editorial a 1186 x 698 y 390 x 844, cinco tipos, filtros, edición, slug automático, revisión/aprobación y formulario estructurado. La ruta real falla cerrada sin identidad. E2E autenticado contra Supabase real queda pendiente."),
    ("Evidencia", "Inventario del build con /biblioteca, /biblioteca/[slug], /guias/[slug], /panel/contenido y seis rutas BFF editoriales; salida Vitest 20/20; inspección visual y mediciones DOM en navegador; registros locales seguros en tmp/codex-content-status. No se cargó contenido real, no se asignaron roles remotos y no se expusieron tokens."),
    ("Reversión", "Mientras la migración no esté aplicada, retirar el código y el archivo SQL restaura el estado previo sin pérdida de datos. Después de aplicarla, exportar contenidos y objetos del bucket antes de una migración de reversión explícita; no eliminar tablas, objetos ni valores enum sin respaldo y confirmación."),
    ("Pendiente / próximo paso", "Aplicar la migración en un entorno Supabase de desarrollo; configurar SUPABASE_CONTENT_BUCKET; asignar cuentas de prueba a community_contributor, academic_editor, coordination y administrator; ejecutar E2E de propiedad, revisión, publicación y Storage real; cargar contenido académico autorizado y solo entonces desplegar web/API."),
]
if len(change_table.rows) != len(rows):
    raise RuntimeError(f"Unexpected change table rows: {len(change_table.rows)}")
for row, (label, value) in zip(change_table.rows, rows):
    row.cells[0].text = label
    row.cells[1].text = value
for cell in change_table.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

checkpoint.cell(0, 0).text = (
    "Siguiente punto de control / 10/08/2026: el sistema dinámico y el estudio editorial RBAC están "
    "implementados y verificados localmente. Próximo: aplicar la migración en Supabase de desarrollo, "
    "configurar content-assets, asignar cuatro cuentas de prueba, completar el E2E autenticado de "
    "carga/revisión/publicación y desplegar solo después de incorporar contenido académico autorizado."
)

history = doc.tables[-1]
if history.cell(0, 0).text != "Versión":
    raise RuntimeError("Version history table not found at expected position")
new_row_xml = deepcopy(history.rows[-1]._tr)
history._tbl.append(new_row_xml)
history_row = history.rows[-1]
history_row.cells[0].text = "2.6"
history_row.cells[1].text = "10/08/2026"
history_row.cells[2].text = (
    "Catálogo y superficies de aprendizaje dinámicos; estudio editorial RBAC para videos, guías, "
    "cuestionarios, flashcards y temas; Storage privado con carga firmada, pruebas 20/20 y QA responsive. "
    "Migración Supabase, E2E autenticado y despliegue pendientes."
)

output.parent.mkdir(parents=True, exist_ok=True)
doc.save(output)
print(output)
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} history_rows={len(history.rows)}")
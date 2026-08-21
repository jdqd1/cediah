from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.shared import RGBColor

root = Path.cwd()
source = root / "CEDIAH Web Status.docx"
output = root / "tmp" / "codex-content-status" / "CEDIAH Web Status.updated-role.docx"
doc = Document(source)

status = next((paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("[~] El estudio de contenido")), None)
if status is None:
    raise RuntimeError("Technical status paragraph not found")
status.text = (
    "[~] El catálogo dinámico, el estudio editorial y la gestión de roles por correo están "
    "implementados localmente; la migración remota, el QA autenticado con Storage real, "
    "las matrículas, los certificados y el despliegue siguen pendientes."
)
status.style = "Checklist"

checkpoint = next(
    (table for table in doc.tables if table.cell(0, 0).text.startswith("Siguiente punto de control")),
    None,
)
if checkpoint is None:
    raise RuntimeError("Checkpoint table not found")
template = next(
    (table for table in doc.tables if any("CHG-006" in cell.text for row in table.rows for cell in row.cells)),
    None,
)
if template is None:
    raise RuntimeError("CHG-006 table not found")
new_tbl_xml = deepcopy(template._tbl)
checkpoint._tbl.addprevious(new_tbl_xml)
change_table = Table(new_tbl_xml, doc._body)
heading = doc.add_paragraph("Ficha de cambio - CHG-007", style="Heading 2")
new_tbl_xml.addprevious(heading._p)
rows = [
    ("Campo", "Contenido"),
    ("ID / fecha / autor", "CHG-007 / 10/08/2026 / Codex con José"),
    ("Fase y objetivo", "Administración segura de roles: bootstrap del primer administrador y asignación/revocación por correo desde la interfaz."),
    ("Qué cambió", "Nueva pantalla /panel/administracion/roles, enlace protegido en el menú y panel; consulta de cuentas Supabase Auth por correo; selector de acción y rol para asignar o revocar."),
    ("Cómo se implementó", "Contratos AdminRole; proveedor Supabase server-side con Auth Admin API y user_roles; rutas Fastify /v1/admin/roles; BFF Next /api/admin/roles; sólo administrator puede mutar roles."),
    ("Archivos / migraciones / servicios", "apps/api/src/providers/supabase-role-management.ts; apps/api/src/app.ts; apps/api/test/admin-roles.test.ts; apps/web/src/components/role-management-screen.tsx; panel/administracion/roles; migración 20260810215000_add_administrator_role_guard.sql."),
    ("Seguridad y privacidad", "La sesión y el rol se validan en cada petición; la clave secreta permanece en Fastify; las cuentas se buscan en Auth; cada cambio se registra en audit_log; trigger y API impiden eliminar el último administrator; no se concede service_role al navegador."),
    ("Validación automática", "PASS: pnpm lint; typecheck de contratos, API y web; pnpm test (5 archivos, 24/24); build API; build Next 16.2.12 en salida aislada con 28 rutas, incluida la administración de roles; git diff --check."),
    ("Validación en navegador", "La ruta /panel/administracion/roles se incluyó en el build y la navegación queda condicionada a roles server-side. QA autenticado contra Supabase real sigue pendiente porque no hay usuarios/roles remotos configurados en este entorno."),
    ("Evidencia", "Salida Vitest 24/24, typecheck, lint, build Next con /api/admin/roles y /panel/administracion/roles, documentación README/architecture/Supabase y pruebas negativas de estudiante/no autenticado."),
    ("Reversión", "Mientras la migración no esté aplicada, retirar el código y el SQL restaura el estado anterior. Después, exportar roles/auditoría antes de una reversión explícita; conservar siempre al menos una cuenta administrator."),
    ("Pendiente / próximo paso", "Aplicar las dos migraciones en Supabase de desarrollo, bootstrappear el primer administrator desde SQL Editor, asignar roles de prueba por correo, completar E2E autenticado y desplegar sólo después de validar el proyecto destino."),
]
if len(change_table.rows) != len(rows):
    raise RuntimeError(f"Unexpected change table rows: {len(change_table.rows)}")
for row, (label, value) in zip(change_table.rows, rows):
    row.cells[0].text = label
    row.cells[1].text = value
for cell in change_table.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

checkpoint.cell(0, 0).text = (
    "Siguiente punto de control / 10/08/2026: el catálogo dinámico, el estudio editorial y la administración "
    "RBAC están implementados y verificados localmente. Próximo: aplicar migraciones en Supabase de desarrollo, "
    "bootstrappear el primer administrator desde SQL Editor, asignar cuentas por correo, completar E2E autenticado "
    "de contenido/Storage/roles y desplegar sólo después de incorporar contenido autorizado."
)

history = doc.tables[-1]
if history.cell(0, 0).text != "Versión":
    raise RuntimeError("Version history table not found")
new_row_xml = deepcopy(history.rows[-1]._tr)
history._tbl.append(new_row_xml)
history_row = history.rows[-1]
history_row.cells[0].text = "2.7"
history_row.cells[1].text = "10/08/2026"
history_row.cells[2].text = (
    "Gestión RBAC por correo desde panel administrador, bootstrap documentado, auditoría y guardia contra "
    "eliminar el último administrator. Tests 24/24 y build con rutas administrativas; migración, E2E y despliegue pendientes."
)

output.parent.mkdir(parents=True, exist_ok=True)
doc.save(output)
print(output)
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} history_rows={len(history.rows)}")
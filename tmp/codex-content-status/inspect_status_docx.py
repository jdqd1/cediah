from pathlib import Path
from docx import Document

path = Path("CEDIAH Web Status.docx")
doc = Document(path)
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
for index, paragraph in enumerate(doc.paragraphs):
    text = " ".join(paragraph.text.split())
    if text:
        print(f"P{index:03d} [{paragraph.style.name}] {text}")
for table_index, table in enumerate(doc.tables):
    print(f"TABLE {table_index}: rows={len(table.rows)} cols={len(table.columns)}")
    for row_index, row in enumerate(table.rows):
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        print(f"  R{row_index:02d}: " + " | ".join(cells))